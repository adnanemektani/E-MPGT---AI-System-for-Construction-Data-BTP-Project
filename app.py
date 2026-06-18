import streamlit as st
from dotenv import load_dotenv
import os, uuid
from datetime import datetime

import fitz
from docx import Document
import openpyxl
from PIL import Image
import imaplib
import email
import io

from pinecone import Pinecone, ServerlessSpec
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain_classic.chains import ConversationalRetrievalChain
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.documents import Document as LCDocument
from langchain_core.retrievers import BaseRetriever
from langchain_core.callbacks import CallbackManagerForRetrieverRun
from typing import List
from htmlTemplates import css, bot_template, user_template
import google.generativeai as genai


class PineconeVectorStore:
    def __init__(self, index, embedding, text_key="text", namespace="default"):
        self._index = index
        self._embedding = embedding
        self._text_key = text_key
        self._namespace = namespace

    def similarity_search(self, query, k=4, namespace=None):
        vec = self._embedding.embed_query(query)
        ns = namespace or self._namespace
        results = self._index.query(vector=vec, top_k=k, include_metadata=True, namespace=ns)
        docs = []
        for match in results.matches:
            text = match.metadata.get(self._text_key, "")
            meta = {k: v for k, v in match.metadata.items() if k != self._text_key}
            docs.append(LCDocument(page_content=text, metadata=meta))
        return docs

    def as_retriever(self, search_kwargs=None):
        ns = (search_kwargs or {}).get("namespace", self._namespace)
        k  = (search_kwargs or {}).get("k", 4)
        store = self

        class _Retriever(BaseRetriever):
            def _get_relevant_documents(self, query: str, *, run_manager: CallbackManagerForRetrieverRun) -> List[LCDocument]:
                return store.similarity_search(query, k=k, namespace=ns)
        return _Retriever()


def extract_text_and_metadata(uploaded_files, project_name, lot_technique, author="User"):
    documents = []
    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
    model = genai.GenerativeModel('gemini-1.5-flash')
    for file in uploaded_files:
        text = ""
        ext = file.name.split('.')[-1].lower()
        if ext == 'pdf':
            doc = fitz.open(stream=file.read(), filetype="pdf")
            for page in doc:
                text += page.get_text()
        elif ext == 'docx':
            doc = Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])
        elif ext == 'xlsx':
            wb = openpyxl.load_workbook(file)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    text += " ".join([str(c) for c in row if c]) + "\n"
        elif ext in ['png', 'jpg', 'jpeg']:
            try:
                img = Image.open(file)
                response = model.generate_content([
                    "Extract all text from this construction document image. "
                    "Also describe any technical diagrams or tables.", img
                ])
                text = response.text
            except Exception as e:
                text = f"Vision error: {e}"
        if text.strip():
            documents.append({
                "content": text,
                "metadata": {
                    "source": file.name,
                    "project": project_name,
                    "lot": lot_technique,
                    "type": ext,
                    "author": author,
                    "date": datetime.now().strftime("%Y-%m-%d")
                }
            })
    return documents


def get_text_chunks(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = []
    for doc in documents:
        for split in splitter.split_text(doc["content"]):
            chunks.append({"text": split, "metadata": doc["metadata"]})
    return chunks


def get_vectorstore(chunks, namespace="default"):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
    index_name = os.getenv("PINECONE_INDEX_NAME")
    if index_name not in [i.name for i in pc.list_indexes()]:
        pc.create_index(
            name=index_name, dimension=384, metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1")
        )
    index = pc.Index(index_name)
    batch = []
    for chunk in chunks:
        vec = embeddings.embed_query(chunk["text"])
        meta = {**chunk["metadata"], "text": chunk["text"]}
        batch.append((str(uuid.uuid4()), vec, meta))
        if len(batch) >= 100:
            index.upsert(vectors=batch, namespace=namespace)
            batch = []
    if batch:
        index.upsert(vectors=batch, namespace=namespace)
    return index


def get_conversation_chain(index, namespace):
    llm = ChatGroq(model_name="llama-3.1-8b-instant")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    retriever = PineconeVectorStore(
        index=index, embedding=embeddings, text_key="text", namespace=namespace
    ).as_retriever(search_kwargs={"namespace": namespace})
    return ConversationalRetrievalChain.from_llm(
        llm=llm, retriever=retriever, return_source_documents=True
    )


def handle_userinput(user_question):
    if st.session_state.conversation is None:
        st.warning("Process a project first!")
        return
    response = st.session_state.conversation({
        'question': user_question,
        'chat_history': st.session_state.chat_history
    })
    st.session_state.chat_history.append(HumanMessage(content=user_question))
    st.session_state.chat_history.append(AIMessage(content=response['answer']))
    conv_id = st.session_state.current_conv_id
    if conv_id and conv_id in st.session_state.conversations:
        if st.session_state.conversations[conv_id]["title"] == "New conversation":
            st.session_state.conversations[conv_id]["title"] = user_question[:40] + "..."
        st.session_state.conversations[conv_id]["history"] = list(st.session_state.chat_history)


def main():
    load_dotenv()
    st.set_page_config(
        page_title="E-MPGT BTP Assistant",
        page_icon="https://i.ibb.co/tpPtBhDK/Whats-App-Image-2026-05-03-at-13-09-45.jpg",
        layout="wide"
    )
    st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)

    # Dark/Light toggle
    st.markdown("""
    <button class="theme-btn" onclick="
        document.body.classList.toggle('light-mode');
        this.textContent = document.body.classList.contains('light-mode') ? 'Dark' : 'Light';
    ">Light</button>
    """, unsafe_allow_html=True)

    # Session state
    for key, val in [
        ("conversation", None), ("chat_history", []),
        ("conversations", {}), ("current_conv_id", None), ("conv_indexes", {})
    ]:
        if key not in st.session_state:
            st.session_state[key] = val

    # ── Sidebar ──────────────────────────────────────────────
    with st.sidebar:
        st.markdown("## E-MPGT Assistant")
        st.markdown("---")

        # Sync Email — fo9
        st.caption("SYNC EMAIL")
        if st.button("Sync Gmail", use_container_width=True):
            with st.spinner("Scanning inbox..."):
                try:
                    project_name_sync = st.session_state.get("_project_name", "Chantier_A")
                    namespace_sync = project_name_sync.lower().replace(" ", "_")
                    lot_sync = st.session_state.get("_lot_tech", "Autre")
                    mail = imaplib.IMAP4_SSL("imap.gmail.com")
                    mail.login(os.getenv("EMAIL_USER"), os.getenv("EMAIL_PASS"))
                    mail.select("inbox")
                    _, messages = mail.search(None, f'(SUBJECT "{project_name_sync}")')
                    new_docs = []
                    for num in messages[0].split():
                        _, data = mail.fetch(num, "(RFC822)")
                        msg = email.message_from_bytes(data[0][1])
                        sender = msg.get("From")
                        for part in msg.walk():
                            if part.get_filename():
                                f = io.BytesIO(part.get_payload(decode=True))
                                f.name = part.get_filename()
                                new_docs.extend(extract_text_and_metadata(
                                    [f], project_name_sync, lot_sync, author=sender
                                ))
                    if new_docs:
                        chunks = get_text_chunks(new_docs)
                        index = get_vectorstore(chunks, namespace=namespace_sync)
                        chain = get_conversation_chain(index, namespace_sync)
                        conv_id = str(uuid.uuid4())
                        st.session_state.conversations[conv_id] = {
                            "title": f"Gmail {project_name_sync} {datetime.now().strftime('%H:%M')}",
                            "history": []
                        }
                        st.session_state.conv_indexes[conv_id] = (index, namespace_sync)
                        st.session_state.current_conv_id = conv_id
                        st.session_state.chat_history = []
                        st.session_state.conversation = chain
                        st.success(f"{len(new_docs)} docs synced!")
                    else:
                        st.info("No new attachments found.")
                    mail.logout()
                except Exception as e:
                    st.error(f"Error: {e}")

        st.markdown("---")

        # Config projet — te7t
        st.caption("CONFIG. PROJET")
        project_name = st.text_input("Nom du Projet", value="Chantier_A", key="pname")
        st.session_state["_project_name"] = project_name
        namespace = project_name.lower().replace(" ", "_")

        lot_tech = st.selectbox("Lot Technique",
            ["Gros Oeuvre", "Electricite", "Plomberie", "VRD", "Autre"], key="ltech")
        st.session_state["_lot_tech"] = lot_tech

        uploaded_files = st.file_uploader(
            "Upload (PDF, Docx, Xlsx, Images)",
            accept_multiple_files=True, key="upfiles"
        )

        if st.button("Process Upload", use_container_width=True, type="primary"):
            with st.spinner("Processing..."):
                docs = extract_text_and_metadata(uploaded_files, project_name, lot_tech)
                chunks = get_text_chunks(docs)
                index = get_vectorstore(chunks, namespace=namespace)
                chain = get_conversation_chain(index, namespace)
                conv_id = str(uuid.uuid4())
                st.session_state.conversations[conv_id] = {
                    "title": f"{project_name} {datetime.now().strftime('%H:%M')}",
                    "history": []
                }
                st.session_state.conv_indexes[conv_id] = (index, namespace)
                st.session_state.current_conv_id = conv_id
                st.session_state.chat_history = []
                st.session_state.conversation = chain
                st.success(f"{project_name} processed!")

        st.markdown("---")

        # New chat + conversations history
        if st.button("New Chat", use_container_width=True):
            conv_id = str(uuid.uuid4())
            st.session_state.conversations[conv_id] = {"title": "New conversation", "history": []}
            st.session_state.current_conv_id = conv_id
            st.session_state.chat_history = []
            st.session_state.conversation = None

        if st.session_state.conversations:
            st.caption("CONVERSATIONS")
            for conv_id, conv in reversed(list(st.session_state.conversations.items())):
                is_active = conv_id == st.session_state.current_conv_id
                if st.button(
                    conv["title"], key=f"conv_{conv_id}",
                    use_container_width=True,
                    type="primary" if is_active else "secondary"
                ):
                    st.session_state.current_conv_id = conv_id
                    st.session_state.chat_history = list(conv["history"])
                    if conv_id in st.session_state.conv_indexes:
                        idx, ns = st.session_state.conv_indexes[conv_id]
                        st.session_state.conversation = get_conversation_chain(idx, ns)

    # ── Chat area ─────────────────────────────────────────────
    if st.session_state.chat_history:
        st.markdown('<div class="chat-container">', unsafe_allow_html=True)
        for message in st.session_state.chat_history:
            if isinstance(message, HumanMessage):
                st.markdown(user_template.replace("{{MSG}}", message.content), unsafe_allow_html=True)
            elif isinstance(message, AIMessage):
                st.markdown(bot_template.replace("{{MSG}}", message.content), unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="welcome">
            <h2>E-MPGT — Intelligent BTP Assistant</h2>
            <p>Upload your construction documents and start asking questions.</p>
        </div>
        """, unsafe_allow_html=True)

    # Fixed input
    def submit():
        q = st.session_state.chat_input
        if q:
            handle_userinput(q)
            st.session_state.chat_input = ""

    st.markdown('<div class="fixed-input-bar">', unsafe_allow_html=True)
    st.text_input(
        "", placeholder="Ask about your construction documents...",
        key="chat_input", on_change=submit, label_visibility="collapsed"
    )
    st.markdown('</div>', unsafe_allow_html=True)


if __name__ == '__main__':
    main()